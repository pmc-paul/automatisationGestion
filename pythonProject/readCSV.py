import csv
import xlsxwriter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

#####################################
# À changer tout le temps
nbOfWeeks = 4  # nb de semaines depuis de debut session pour la moyenne des heures
nbTimeCoursLastWeek = 1  #Le temps à ajouter si on a un cours ou wtv (ex. le 1h de présentation technique avec les prof)
nbTimeCoursNextWeek = 1  #same as nbTimeCoursLastWeek, mais pour la semaine prochaine

sprintLastWeek = 'S6-2' #sprint semaine passée
lastSprintWeek = '1' + ',' #semaine de sprint de la semaine passé  (NE pas effacer)

sprintNextWeek = 'S6-2' #sprint de la semaine prochaine
nextSprintWeek = '2' + ',' #semaine de sprint de la semaine prochaine

# Noms des fichiers à import
# fichier raodmap = roadmap.csv
# fichier time Log = timeLog.csv
# word tableau de bord = tableauBord.csv

#####################################


# [name, nb heures travaillées, tâches effectuées, moyenne/semaine heures travaillées, prévision heures]
name = 0; hours = 1; tasksDone = 2; totalAvergHours = 3; previsionHours = 4; tasksToDo = 5;
members = [['Krystel Smith', 0, '', 0, 0, ''],
           ['William Babin-Demers', 0, '', 0, 0, ''],
           ['Victoria Pitz Clairoux', 0, '', 0, 0, ''],
           ['Olivier Roy', 0, '', 0, 0, ''],
           ['Marc-Olivier Thibault', 0, '', 0, 0, ''],
           ['Philippe Boulet', 0, '', 0, 0, ''],
           ['Alexandre Filion', 0, '', 0, 0, '']]

modules = [['Navigation', 0], ['Application Web', 0], ['Plateforme mobile', 0], ['Bras robotisé', 0], ['Gestion', 0]]
# Tâche,Personne,Nb heures travaillées,Module/Type,Date,Sprint,Précision type

def readTimelog():
    taskCol = 0; nameCol = 1; hourCol = 2; sprintCol = 5;  #pour naviger dans les colonnes du csv
    with open('timeLog.csv') as csv_file:
        csv_reader = csv.reader(csv_file, delimiter=',')
        for idx, row in enumerate(csv_reader):
            if idx > 0:  #skip la première ligne avec les titres des sections
                for member in members:
                    if member[name] in row[nameCol]:
                        if lastSprintWeek in row[sprintCol] and sprintLastWeek in row[sprintCol]:
                            member[hours] += (float(row[hourCol]))
                            member[tasksDone] += '- ' + (row[taskCol] + ' (' + row[hourCol] + 'h) ')
                        else:
                            if row[hourCol] != '':
                                member[totalAvergHours] += float(row[hourCol])
        for idx, member in enumerate(members):
            member[hours] += nbTimeCoursLastWeek
            member[totalAvergHours] = (member[totalAvergHours] + member[hours]) / nbOfWeeks

def readRoadmap():
    taskCol = 0; roadSprintCol = 1; roadModuleCol = 2; AssigneeCol = 3; estiHoursCol = 5; stageCol = 9 #pour naviger dans les colonnes du csv

    with open('roadmap.csv') as csv_file:
        csv_reader = csv.reader(csv_file, delimiter=',')
        for idx, row in enumerate(csv_reader):
            if idx > 0:
                if nextSprintWeek in row[roadSprintCol] and sprintNextWeek in row[roadSprintCol]:
                    for member in members:
                        if member[name] in row[AssigneeCol] and row[stageCol] != 'Done':
                            nbPersTache = 1;
                            nbPersTache += row[AssigneeCol].count(',')
                            nbHoursTask = float(row[estiHoursCol])/nbPersTache
                            member[tasksToDo] += '- ' + (row[taskCol] + ' (' + "{:g}".format(nbHoursTask)  + 'h) ')
                            if row[estiHoursCol] != '':
                                member[previsionHours] += float(row[estiHoursCol])/nbPersTache
                            for module in modules:
                                if module[0] in row[roadModuleCol] and row[estiHoursCol] != '':
                                    module[1] += float(row[estiHoursCol])/nbPersTache
        for member in members:
            member[previsionHours] += nbTimeCoursNextWeek
        for module in modules:
            print(module[0], '', module[1])   #temps par module



# ##############  WRITE EXCEL
def writeExcel():
    workbook = xlsxwriter.Workbook('tableTimeLog.xlsx')
    worksheet = workbook.add_worksheet()
    bold = workbook.add_format({'bold': 1})

    # Add the worksheet data that the charts will refer to.
    headings = ['Name', 'Nb heures cette semaine', 'Moyenne par semaine', 'Prévision', 'Tâches faites', 'Tâche semaine pro']

    worksheet.write_row('A1', headings, bold)
    for index, membeer in enumerate(members):
        worksheet.write_row('A' + str(index + 2), [membeer[index] for index in [name, hours, totalAvergHours, previsionHours, tasksDone, tasksToDo]])

    chart1 = workbook.add_chart({'type': 'column'})

    # Add a chart title and some axis labels.
    chart1.set_title ({'name': 'Gestion du temps'})
    chart1.set_y_axis({'name': "Nb d'heures"})

    chart1.add_series({
        'name':       '=Sheet1!$B$1',
        'categories': '=Sheet1!$A$2:$A$8',
        'values':     '=Sheet1!$B$2:$B$8',
    })

    # Configure a second series. Note use of alternative syntax to define ranges. (2 méthodes possibles)
    # Or using a list of values instead of category/value formulas:  [sheetname, first_row, first_col, last_row, last_col]

    chart1.add_series({
        'name':       ['Sheet1', 0, 3],
        'categories': ['Sheet1', 1, 0, 7, 0],
        'values':     ['Sheet1', 1, 3, 7, 3],
    })
    # Create a new column chart. This will use this as the secondary chart.
    line_chart1 = workbook.add_chart({'type': 'line'})

    # Configure the data series for the secondary chart.
    line_chart1.add_series({
        'name': ['Sheet1', 0, 2],
        'categories': ['Sheet1', 1, 0, 7, 0],
        'values': ['Sheet1', 1, 2, 7, 2],
    })

    # Combine the charts.
    chart1.combine(line_chart1)

    # Set an Excel chart style.
    chart1.set_style(10)
    chart1.set_legend({'position': 'bottom'})
    # Insert the chart into the worksheet (with an offset).
    worksheet.insert_chart('G12', chart1, {'x_offset': 25, 'y_offset': 10})


    # task
    # Add the worksheet data that the charts will refer to.
    headings = ['Name', 'Nb heures cette semaine', 'Moyenne par semaine', 'Prévision']

    worksheet.write_row('A1', headings, bold)
    for index, membeer in enumerate(members):
        worksheet.write_row('A' + str(index + 2), [membeer[index] for index in [name, hours, totalAvergHours, previsionHours]])

    workbook.close()

# #############   Change word
def changeWord():
    doc = Document('tableauBord.docx')
    # modification du tableau des tâches
    for member in members:
        for x in range(1, 8):
            if member[0] in doc.tables[0].cell(x, 1).text:
                for colTableau, taskToWrite in zip((2,3), (tasksDone, tasksToDo)):
                    doc.tables[0].cell(x, colTableau).text = ''
                    for idx, letter in enumerate(member[taskToWrite]):
                        paragraph = doc.tables[0].cell(x, colTableau).paragraphs[0];
                        if idx > 0 and letter != '-' :
                            font = paragraph.add_run(letter).font
                        elif idx > 0 and letter == '-':
                            paragraph.add_run().add_break(WD_BREAK.LINE)
                            font = paragraph.add_run(letter).font
                        else:
                            font = paragraph.add_run(letter).font
                        font.size = Pt(10)
                        font.name = 'Times New Roman'
    doc.save("newTableauBord.docx")



readTimelog()
readRoadmap()
writeExcel()
changeWord()
